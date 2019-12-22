from docx import Document
from docx.shared import Inches

def max_length(*lists):
    return max(list(map(lambda l: len(l), *lists)))


def my_zip(*lists):
    ml = max_length(lists)
    for l in lists:
        r = ml - len(l)
        if r != 0:
            for __ in range(r):
                l.append(None)
    return tuple(map(lambda *x: x, *lists))


l = ['mptu', 2, 3, 9]
ll = ['yare', 7, 6, 5, 4]
lll = ['fff', 10, 20, '30tri', 40, 50]


def task2():
    print(str(my_zip(l, ll, lll)))


def readTexFile(path):
    texdoc = []
    with open(path) as fin:
        for line in fin:
            texdoc.append(line.replace('width=.5\\textwidth', 'width=.9\\textwidth'))
    return texdoc


def writeTexFile(texdoc, path):
    with open(path, 'w') as fout:
        for i in range(len(texdoc)):
            fout.write(texdoc[i])

def conjecturer(strList):
    return list(map(lambda s: s.replace("hypothesis", "conjecture").replace("hypotheses", "conjectures"), strList))

def increaseFontsize(docClassPath, newDocClassPath):
    strList = readTexFile(docClassPath)
    c = list(map(lambda s: s.replace("", "") if s.count("\\fontsize") != 0 else s, strList))
    writeTexFile(c, newDocClassPath)

def notCommentedSection(line):
    if(line.count("subsection") != 0 or line.count("\\section") != 0):
        for c in (0, line.index("section")):
            if c == "%":
                return False
            else:
                return True
    else:
        return False

def docxStructure(strList):
    structured = []
    tostruc = list(filter(lambda s: notCommentedSection(s), strList))
    for line in tostruc:
        k = line.count("sub")
        structured.append(("\t" * k) + line[line.index('{')+1:line.index('}')])
    return structured

def texStructure(strList):
    structured = []
    tostruc = list(filter(lambda s: notCommentedSection(s), strList))
    for line in tostruc:
        structured.append(line[line.index('\\'):line.index('}') + 1] + "\n")
    return structured

def createDocx(list):
    document = Document()
    document.add_heading('Structure', 0)
    for line in list:
        document.add_paragraph(line)
    document.save('HW4.docx')

def getStyle(doc):
    return doc[0 : doc.index("\\begin{document}\n") + 1]


def tex():
    print("Please input path to the *.tex file")
    path = str(input())

    print("Please input filename with .tex ")
    filename = str(input())
    doc = readTexFile(path + filename)
    a = conjecturer(doc)
    writeTexFile(a, path + 'conjectured.tex')
    createDocx(docxStructure(doc))
    struc = getStyle(doc)
    struc.extend(texStructure(doc))
    struc.append("\\end{document}")
    writeTexFile(struc, path + 'structure.tex')